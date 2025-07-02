﻿import io
import xmlrpc.client
import streamlit as st
from docx import Document
import datetime
from typing import Tuple, Optional, List, Dict, Any
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import os

# Template configuration - using absolute paths from o3_templates.py
TEMPLATE_OPTIONS = {
    "employment_letter_arabic": {
        "name": "Employment letter - Arabic",
        "path": r"C:\Users\Geeks\source\repos\CSTemplates\Employment Letter - ARABIC.docx",
        "description": "Employment letter in Arabic"
    },
    "employment_letter": {
        "name": "Employment letter", 
        "path": r"C:\Users\Geeks\source\repos\CSTemplates\Employment Letter .docx",  # Note the space before .docx
        "description": "Standard employment letter in English"
    },
    "employment_letter_embassy": {
        "name": "Employment letter to embassies",
        "path": r"C:\Users\Geeks\source\repos\CSTemplates\Employment Letter to Embassies.docx", 
        "description": "Employment letter for visa/embassy purposes"
    },
    "experience_letter": {
        "name": "Experience letter",
        "path": r"C:\Users\Geeks\source\repos\CSTemplates\Experience Letter.docx",
        "description": "Experience certificate for former employees"
    }
}

# List of countries for embassy letters
COUNTRIES = [
    "Afghanistan", "Albania", "Algeria", "Andorra", "Angola", "Antigua and Barbuda",
    "Argentina", "Armenia", "Australia", "Austria", "Azerbaijan", "Bahamas", "Bahrain",
    "Bangladesh", "Barbados", "Belarus", "Belgium", "Belize", "Benin", "Bhutan", "Bolivia",
    "Bosnia and Herzegovina", "Botswana", "Brazil", "Brunei", "Bulgaria", "Burkina Faso",
    "Burundi", "Cabo Verde", "Cambodia", "Cameroon", "Canada", "Central African Republic",
    "Chad", "Chile", "China", "Colombia", "Comoros", "Congo", "Costa Rica",
    "Croatia", "Cuba", "Cyprus", "Czech Republic", "Democratic Republic of the Congo",
    "Denmark", "Djibouti", "Dominica", "Dominican Republic", "Ecuador", "Egypt", "El Salvador",
    "Equatorial Guinea", "Eritrea", "Estonia", "Ethiopia", "Fiji", "Finland", "France", 
    "Gabon", "Gambia", "Georgia", "Germany", "Ghana", "Greece", "Grenada", "Guatemala",
    "Guinea", "Guinea-Bissau", "Guyana", "Haiti", "Honduras", "Hungary", "Iceland", "India",
    "Indonesia", "Iran", "Iraq", "Ireland", "Italy", "Jamaica", "Japan", "Jordan",
    "Kazakhstan", "Kenya", "Kiribati", "Kuwait", "Kyrgyzstan", "Laos", "Latvia", "Lebanon",
    "Lesotho", "Liberia", "Libya", "Liechtenstein", "Lithuania", "Luxembourg", "Madagascar",
    "Malawi", "Malaysia", "Maldives", "Mali", "Malta", "Marshall Islands", "Mauritania",
    "Mauritius", "Mexico", "Micronesia", "Moldova", "Monaco", "Mongolia", "Montenegro",
    "Morocco", "Mozambique", "Myanmar", "Namibia", "Nauru", "Nepal", "Netherlands",
    "New Zealand", "Nicaragua", "Niger", "Nigeria", "North Korea", "North Macedonia",
    "Norway", "Oman", "Pakistan", "Palau", "Palestine", "Panama", "Papua New Guinea",
    "Paraguay", "Peru", "Philippines", "Poland", "Portugal", "Qatar", "Romania", "Russia",
    "Rwanda", "Saint Kitts and Nevis", "Saint Lucia", "Saint Vincent and the Grenadines",
    "Samoa", "San Marino", "Sao Tome and Principe", "Saudi Arabia", "Senegal", "Serbia",
    "Seychelles", "Sierra Leone", "Singapore", "Slovakia", "Slovenia", "Solomon Islands",
    "Somalia", "South Africa", "South Korea", "South Sudan", "Spain", "Sri Lanka", "Sudan",
    "Suriname", "Sweden", "Switzerland", "Syria", "Tajikistan", "Tanzania", "Thailand",
    "Timor-Leste", "Togo", "Tonga", "Trinidad and Tobago", "Tunisia", "Turkey",
    "Turkmenistan", "Tuvalu", "Uganda", "Ukraine", "United Arab Emirates", "United Kingdom",
    "United States", "Uruguay", "Uzbekistan", "Vanuatu", "Venezuela", "Vietnam", "Yemen",
    "Zambia", "Zimbabwe"
]

def get_arabic_name(employee: Dict[str, Any]) -> str:
    """Get the Arabic name of the employee"""
    name = employee.get("x_studio_employee_arabic_name", "").strip()
    if name:
        return name
    return employee.get("name", "").strip()

def get_partner_address(models: xmlrpc.client.ServerProxy, uid: int, db: str, password: str, partner_id: int) -> str:
    """Get the partner's address"""
    try:
        fields = ["street", "street2", "city", "zip", "country_id"]
        result = models.execute_kw(db, uid, password, "res.partner", "read", [[partner_id]], {"fields": fields})
        if result:
            partner = result[0]
            country = ""
            if partner.get("country_id"):
                if isinstance(partner["country_id"], (list, tuple)) and len(partner["country_id"]) > 1:
                    country = partner["country_id"][1]
                else:
                    country = str(partner["country_id"])
            address_parts = [partner.get("street", ""), partner.get("street2", ""), 
                           partner.get("city", ""), partner.get("zip", ""), country]
            address_str = ", ".join([part for part in address_parts if part])
            return address_str
        return ""
    except Exception as e:
        return ""

def get_arabic_partner_address(models: xmlrpc.client.ServerProxy, uid: int, db: str, password: str, partner_id: int) -> str:
    """Get the Arabic address of the partner"""
    try:
        fields = ["x_studio_arabic_address"]
        result = models.execute_kw(db, uid, password, "res.partner", "read", [[partner_id]], {"fields": fields})
        if result:
            partner = result[0]
            arabic_address = partner.get("x_studio_arabic_address", "")
            return arabic_address.strip() if arabic_address else ""
        return ""
    except Exception:
        return ""

def get_company_registrar(models: xmlrpc.client.ServerProxy, uid: int, db: str, password: str, company_id: int) -> str:
    """Get the company registrar number"""
    try:
        result = models.execute_kw(db, uid, password, "res.company", "read", [[company_id]], {"fields": ["company_registry"]})
        if result:
            return result[0].get("company_registry", "")
        return ""
    except Exception:
        return ""

def get_company_arabic_name(models: xmlrpc.client.ServerProxy, uid: int, db: str, password: str, company_id: int) -> str:
    """Get the Arabic name of the company"""
    try:
        result = models.execute_kw(db, uid, password, "res.company", "read", [[company_id]], {"fields": ["arabic_name"]})
        if result:
            return result[0].get("arabic_name", "")
        return ""
    except Exception:
        return ""

def get_head_people_and_culture(models: xmlrpc.client.ServerProxy, uid: int, db: str, password: str, company_id: int) -> str:
    """Get the head of people and culture"""
    try:
        domain = [('company_id', '=', company_id), ('job_id.name', 'ilike', 'head of people and culture')]
        head_ids = models.execute_kw(db, uid, password, 'hr.employee', 'search', [domain])
        if head_ids:
            head_data = models.execute_kw(db, uid, password, 'hr.employee', 'read', [head_ids[0]], {'fields': ['name']})
            if head_data:
                return head_data[0].get('name', 'Faisal Abdullah AlMamun')
        return "Faisal Abdullah AlMamun"  # Default value
    except Exception:
        return "Faisal Abdullah AlMamun"

def get_head_people_and_culture_arabic(models: xmlrpc.client.ServerProxy, uid: int, db: str, password: str, company_id: int) -> str:
    """Get the Arabic name of head of people and culture"""
    try:
        domain = [('company_id', '=', company_id), ('job_id.name', 'ilike', 'head of people and culture')]
        head_ids = models.execute_kw(db, uid, password, 'hr.employee', 'search', [domain])
        if head_ids:
            head_data = models.execute_kw(db, uid, password, 'hr.employee', 'read', [head_ids[0]], {'fields': ['x_studio_employee_arabic_name', 'name']})
            if head_data:
                return get_arabic_name(head_data[0]) or "فيصل عبدالله المأمون"
        return "فيصل عبدالله المأمون"  # Default value
    except Exception:
        return "فيصل عبدالله المأمون"

def derive_country_from_address(address: str) -> str:
    """Derive country from address string"""
    if not address:
        return ""
    if "\n" in address:
        lines = [line.strip() for line in address.split("\n") if line.strip()]
        if lines:
            return lines[-1]
    parts = [part.strip() for part in address.split(",") if part.strip()]
    if parts:
        return parts[-1]
    return ""

def enrich_employee_data_for_template(employee_data: Dict[str, Any], models: xmlrpc.client.ServerProxy, 
                                    uid: int, db: str, password: str) -> Dict[str, Any]:
    """Enrich employee data with additional information needed for templates"""
    enriched_data = employee_data.copy()
    
    # Ensure we have the employee ID
    employee_id = employee_data.get('id')
    if not employee_id:
        return enriched_data
    
    # Get additional fields that might not be in the basic employee data
    try:
        additional_fields = [
            'x_studio_joining_date',
            'x_studio_contract_end_date', 
            'x_studio_employee_arabic_name',
            'identification_id'
        ]
        
        additional_data = models.execute_kw(
            db, uid, password, 'hr.employee', 'read',
            [employee_id], {'fields': additional_fields}
        )
        
        if additional_data:
            employee_data.update(additional_data[0])
    except:
        pass
    
    # Get wage from contracts if available
    try:
        contracts = models.execute_kw(db, uid, password, 'hr.contract', 'search_read',
                                    [[('employee_id', '=', employee_id)]],
                                    {'fields': ['wage'], 'limit': 1})
        wage = contracts[0].get('wage', 0.0) if contracts else 0.0
        enriched_data['wage'] = wage
    except:
        enriched_data['wage'] = 0.0
    
    # Process joining date
    joining_date_raw = employee_data.get('x_studio_joining_date', '')
    if joining_date_raw:
        try:
            joining_date_dt = datetime.datetime.strptime(joining_date_raw, "%Y-%m-%d")
            enriched_data['joining_date'] = joining_date_dt.strftime("%d/%m/%Y")
        except:
            enriched_data['joining_date'] = joining_date_raw
    else:
        enriched_data['joining_date'] = ''
    
    # Process contract end date
    contract_end_date_raw = employee_data.get('x_studio_contract_end_date', '')
    if contract_end_date_raw:
        try:
            contract_end_date_dt = datetime.datetime.strptime(contract_end_date_raw, "%Y-%m-%d")
            enriched_data['contract_end_date'] = contract_end_date_dt.strftime("%d/%m/%Y")
        except:
            enriched_data['contract_end_date'] = contract_end_date_raw
    else:
        enriched_data['contract_end_date'] = ''
    
    # Get department
    department_field = employee_data.get("department_id")
    if department_field and isinstance(department_field, (list, tuple)):
        enriched_data['department'] = department_field[1] if len(department_field) > 1 else str(department_field[0])
    else:
        enriched_data['department'] = str(department_field) if department_field else ''
    
    # Get Arabic name
    enriched_data['arabic_name'] = get_arabic_name(employee_data)
    
    # Get company information
    company_field = employee_data.get("company_id")
    if company_field and isinstance(company_field, (list, tuple)) and len(company_field) > 0:
        company_id = company_field[0]
        enriched_data['company'] = company_field[1] if len(company_field) > 1 else str(company_field[0])
        enriched_data['company_registrar'] = get_company_registrar(models, uid, db, password, company_id)
        enriched_data['company_arabic_name'] = get_company_arabic_name(models, uid, db, password, company_id) or enriched_data['company']
        enriched_data['head_people_culture'] = get_head_people_and_culture(models, uid, db, password, company_id)
        enriched_data['head_people_culture_arabic'] = get_head_people_and_culture_arabic(models, uid, db, password, company_id)
    else:
        enriched_data['company'] = str(company_field) if company_field else ''
        enriched_data['company_registrar'] = ''
        enriched_data['company_arabic_name'] = ''
        enriched_data['head_people_culture'] = "Faisal Abdullah AlMamun"
        enriched_data['head_people_culture_arabic'] = "فيصل عبدالله المأمون"
    
    # Get work address
    address_field = employee_data.get("address_id")
    if address_field and isinstance(address_field, (list, tuple)) and len(address_field) > 0:
        partner_id = address_field[0]
        enriched_data['work_address'] = get_partner_address(models, uid, db, password, partner_id)
        enriched_data['arabic_work_address'] = get_arabic_partner_address(models, uid, db, password, partner_id)
    else:
        enriched_data['work_address'] = str(address_field) if address_field else ''
        enriched_data['arabic_work_address'] = ''
    
    enriched_data['company_country'] = derive_country_from_address(enriched_data['work_address'])
    enriched_data['first_name'] = employee_data.get('name', '').split()[0] if employee_data.get('name') else ''
    enriched_data['identification'] = employee_data.get('identification_id', '').strip()
    
    return enriched_data

def replace_placeholder_in_paragraph(paragraph, placeholder, replacement):
    """Replace placeholder in paragraph while preserving formatting"""
    replaced = False
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            replaced = True
    if not replaced:
        full_text = "".join(run.text for run in paragraph.runs)
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, replacement)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_text

def remove_empty_paragraphs(doc: Document) -> None:
    """Remove empty paragraphs from document"""
    for para in list(doc.paragraphs):
        if not para.text.strip():
            p = para._element
            p.getparent().remove(p)

def fill_template(template_path: str, employee_data: Dict[str, Any], is_arabic: bool = False) -> Optional[bytes]:
    """Fill template with employee data"""
    # Initialize debug_info if not exists
    if 'debug_info' not in st.session_state:
        st.session_state.debug_info = {}
    
    if not os.path.exists(template_path):
        # Log error for debugging
        st.session_state.debug_info['template_error'] = f"Template file not found: {template_path}"
        return None
    
    try:
        doc = Document(template_path)
    except Exception as e:
        st.session_state.debug_info['template_error'] = f"Error loading document: {str(e)}"
        return None
    
    current_date = datetime.date.today().strftime("%d/%m/%Y")
    
    # Define placeholders
    placeholders = {
        "(Current Date)": current_date,
        "(First and Last Name)": employee_data.get('name', ''),
        "(First Name)": employee_data.get('first_name', ''),
        "(Position)": employee_data.get('job_title', ''),
        "(Salary)": str(employee_data.get('wage', 0)),
        "(DD/MM/YYYY)": employee_data.get('joining_date', ''),
        "(Country)": employee_data.get('country', ''),
        "(Start Date)": employee_data.get('start_date', ''),
        "(End Date)": employee_data.get('end_date', ''),
        "(Company)": employee_data.get('company', ''),
        "(Work address)": employee_data.get('work_address', ''),
        "(Work Address)": employee_data.get('work_address', ''),
        "(Arabic Work address)": employee_data.get('arabic_work_address', ''),
        "(CR)": employee_data.get('company_registrar', ''),
        "(Company Country)": employee_data.get('company_country', ''),
        "(CompanyA)": employee_data.get('company_arabic_name', ''),
        "(P&C)": employee_data.get('head_people_culture', ''),
        "(AP&C)": employee_data.get('head_people_culture_arabic', ''),
        "(الاسم الكامل)": employee_data.get("arabic_name", employee_data.get('name', '')) if is_arabic else employee_data.get('name', ''),
        "(بلد الوجهة)": employee_data.get('country', ''),
        "(تاريخ البداية)": employee_data.get('start_date', ''),
        "(تاريخ النهاية)": employee_data.get('end_date', ''),
        "(Contract End Date)": employee_data.get('contract_end_date', ''),
        "(Department)": employee_data.get('department', '')
    }
    
    placeholders = {k: str(v) for k, v in placeholders.items()}
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_placeholder_in_paragraph(para, key, value)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():
                        replace_placeholder_in_paragraph(para, key, value)
    
    # Replace in headers
    for section in doc.sections:
        for para in section.header.paragraphs:
            for key, value in placeholders.items():
                replace_placeholder_in_paragraph(para, key, value)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in placeholders.items():
                            replace_placeholder_in_paragraph(para, key, value)
    
    # Replace in footers with smaller font
    for section in doc.sections:
        for para in section.footer.paragraphs:
            text = para.text
            for key, value in placeholders.items():
                if key in text:
                    text = text.replace(key, value)
            para.text = text
            for run in para.runs:
                run.font.size = Pt(8)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        for key, value in placeholders.items():
                            if key in text:
                                text = text.replace(key, value)
                        para.text = text
                        for run in para.runs:
                            run.font.size = Pt(8)
    
    remove_empty_paragraphs(doc)
    
    output_stream = io.BytesIO()
    doc.save(output_stream)
    return output_stream.getvalue()

def detect_template_intent(query: str) -> Optional[str]:
    """Detect if user is asking for a template"""
    query_lower = query.lower()
    
    # Keywords for template requests
    template_keywords = [
        'employment letter', 'employment certificate', 'work certificate',
        'experience letter', 'experience certificate', 'service certificate',
        'embassy letter', 'visa letter', 'travel letter',
        'arabic letter', 'letter in arabic', 'arabic certificate',
        'template', 'document', 'certificate', 'letter'
    ]
    
    # Check for specific template types
    if any(keyword in query_lower for keyword in ['arabic', 'عربي', 'بالعربية']):
        if any(keyword in query_lower for keyword in ['employment', 'work', 'job']):
            return 'employment_letter_arabic'
    
    if any(keyword in query_lower for keyword in ['embassy', 'visa', 'travel', 'consulate']):
        return 'employment_letter_embassy'
    
    if any(keyword in query_lower for keyword in ['experience', 'service', 'former', 'past']):
        return 'experience_letter'
    
    if any(keyword in query_lower for keyword in ['employment letter', 'work certificate', 'employment certificate']):
        return 'employment_letter'
    
    # Check if any template keyword is present
    if any(keyword in query_lower for keyword in template_keywords):
        return 'general_template_request'
    
    return None

def parse_embassy_details(query: str) -> Dict[str, Any]:
    """Parse embassy letter details from query"""
    details = {
        'country': None,
        'start_date': None,
        'end_date': None
    }
    
    query_lower = query.lower()
    
    # Try to find country
    for country in COUNTRIES:
        if country.lower() in query_lower:
            details['country'] = country
            break
    
    # Try to parse dates
    import re
    from datetime import datetime, timedelta
    
    # Look for date patterns
    date_pattern = r'(\d{1,2})[/\-](\d{1,2})(?:[/\-](\d{2,4}))?'
    dates = re.findall(date_pattern, query)
    
    if dates:
        current_year = datetime.now().year
        parsed_dates = []
        
        for date_match in dates:
            month, day, year = date_match
            if not year:
                year = current_year
            elif len(year) == 2:
                year = 2000 + int(year)
            else:
                year = int(year)
            
            try:
                parsed_date = datetime(year, int(month), int(day))
                parsed_dates.append(parsed_date)
            except:
                continue
        
        if parsed_dates:
            details['start_date'] = min(parsed_dates).strftime('%Y-%m-%d')
            details['end_date'] = max(parsed_dates).strftime('%Y-%m-%d')
    
    return details

def generate_template(template_type: str, employee_data: Dict[str, Any], 
                     embassy_details: Optional[Dict[str, Any]] = None) -> Optional[Tuple[bytes, str]]:
    """Generate a template for the employee"""
    
    # Initialize debug_info at the very beginning of this function
    if 'debug_info' not in st.session_state:
        st.session_state.debug_info = {}
    
    # Debug logging
    st.session_state.debug_info['template_generation'] = {
        'template_type': template_type,
        'employee_name': employee_data.get('name'),
        'embassy_details': embassy_details
    }
    
    # Get Odoo connection from session state
    if not hasattr(st.session_state, 'odoo_models') or not st.session_state.odoo_models:
        st.session_state.debug_info['template_generation']['error'] = 'No Odoo connection'
        return None
    
    models = st.session_state.odoo_models
    uid = st.session_state.odoo_uid
    db = st.session_state.db
    password = st.session_state.password
    
    # Enrich employee data
    enriched_data = enrich_employee_data_for_template(
        employee_data, models, uid, db, password
    )
    
    # Add embassy details if applicable
    if template_type == 'employment_letter_embassy' and embassy_details:
        enriched_data['country'] = embassy_details.get('country', '')
        if embassy_details.get('start_date'):
            try:
                start_dt = datetime.datetime.strptime(embassy_details['start_date'], '%Y-%m-%d')
                enriched_data['start_date'] = start_dt.strftime('%d/%m/%Y')
            except:
                enriched_data['start_date'] = embassy_details.get('start_date', '')
        else:
            enriched_data['start_date'] = ''
            
        if embassy_details.get('end_date'):
            try:
                end_dt = datetime.datetime.strptime(embassy_details['end_date'], '%Y-%m-%d')
                enriched_data['end_date'] = end_dt.strftime('%d/%m/%Y')
            except:
                enriched_data['end_date'] = embassy_details.get('end_date', '')
        else:
            enriched_data['end_date'] = ''
    else:
        enriched_data['country'] = ''
        enriched_data['start_date'] = ''
        enriched_data['end_date'] = ''
    
    # Get template info
    template_info = TEMPLATE_OPTIONS.get(template_type)
    if not template_info:
        st.session_state.debug_info['template_generation']['error'] = f'Template type not found: {template_type}'
        return None
    
    template_path = template_info['path']
    is_arabic = template_type == 'employment_letter_arabic'
    
    st.session_state.debug_info['template_generation']['template_path'] = template_path
    st.session_state.debug_info['template_generation']['enriched_data'] = enriched_data
    
    # Generate document
    doc_bytes = fill_template(template_path, enriched_data, is_arabic)
    if not doc_bytes:
        return None
    
    # Create filename
    filename = f"{template_info['name']} - {enriched_data['name']}.docx"
    
    st.session_state.debug_info['template_generation']['success'] = True
    st.session_state.debug_info['template_generation']['filename'] = filename
    
    return doc_bytes, filename